from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output"
DOCX_PATH = OUTPUT_DIR / "romanian_water_infrastructure_feasibility_study.docx"


PALETTE = {
    "navy": RGBColor(28, 50, 67),
    "teal": RGBColor(40, 104, 111),
    "sage": RGBColor(109, 143, 117),
    "light_teal": "DCEBEB",
    "light_sage": "E8EFE9",
    "light_blue": "EAF2F7",
    "light_red": "F7E8E8",
    "light_gold": "F6F0DD",
    "border": "AAB7BE",
    "dark": "1C3243",
}


def load_json(path: Path) -> dict:
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {}


def set_cell_shading(cell, fill: str) -> None:
    if isinstance(fill, RGBColor):
        fill = str(fill)
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = PALETTE["border"], size: str = "6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm: float) -> None:
    width = int(width_cm * 567)
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths_cm) * 567)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width_cm in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width_cm * 567)))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, width_cm in enumerate(widths_cm):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width_cm)


def style_table(table, header_fill=PALETTE["dark"], band_fill="F6F8FA") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(8.6)
        if r_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
        elif r_idx % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, band_fill)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_document_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.2)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 25, PALETTE["navy"]),
        ("Heading 1", 17, PALETTE["navy"]),
        ("Heading 2", 13, PALETTE["teal"]),
        ("Heading 3", 11.2, PALETTE["sage"]),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name in {"Title", "Heading 1"} else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.paragraph_format.space_before = Pt(10 if name == "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(5)

    caption = styles["Caption"]
    caption.font.name = "Aptos"
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(84, 97, 108)


def add_cover(doc: Document, report_inputs: dict) -> None:
    loc = report_inputs.get("location", {})
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("DRAFT FEASIBILITY RESEARCH REPORT")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = PALETTE["teal"]

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Legal and Technical Feasibility Study\n")
    sub = title.add_run("Canal, Dam, or Watercourse Diversion for Community / Farm Water Supply")
    sub.font.size = Pt(18)
    sub.font.color.rgb = PALETTE["teal"]

    doc.add_paragraph()
    fields = [
        ("Jurisdiction", "Romania - water-management regulation by A.N. Apele Romane, ABA, and SGA"),
        ("Study area", f"AOI centered near lat {loc.get('lat', 'TBD')}, lon {loc.get('lon', 'TBD')} / size {loc.get('size_km', 'TBD')} km"),
        ("Project type", "Potential canal, small reservoir/dam, water intake, or hydraulic diversion"),
        ("Primary decision", "Preliminary legal feasibility only; formal design requires certified studies and authority review"),
        ("Prepared", date.today().isoformat()),
    ]
    for k, v in fields:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.2)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{k}: ")
        r.bold = True
        r.font.color.rgb = PALETTE["teal"]
        p.add_run(v)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.paragraph_format.left_indent = Cm(0.2)
    note.paragraph_format.right_indent = Cm(0.2)
    note.paragraph_format.space_before = Pt(12)
    note.paragraph_format.space_after = Pt(12)
    note_run = note.add_run("Important use note: ")
    note_run.bold = True
    note_run.font.color.rgb = PALETTE["navy"]
    note.add_run(
        "This report is a research-grade feasibility model and document template. "
        "It identifies the legal and technical checks that should be satisfied before promoting a water "
        "infrastructure project, but it is not a substitute for a documentatie tehnica prepared by an "
        "entity attested by the competent water authority, nor for legal advice from Romanian counsel."
    )

    doc.add_page_break()


def add_static_toc(doc: Document) -> None:
    doc.add_heading("Contents", level=1)
    items = [
        "1. Executive conclusion",
        "2. Project definition and data basis",
        "3. Romanian legal framework",
        "4. Apele Romane permitting roadmap",
        "5. Baseline technical conditions",
        "6. Alternatives analysis",
        "7. Impact assessment and mitigation",
        "8. Compliance matrix",
        "9. Minimum documentation package",
        "10. Conclusions and next actions",
        "Appendix A. Legal source notes",
        "Appendix B. Technical checklist",
    ]
    for item in items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(2)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(2)


def add_matrix_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    for row_values in rows:
        title = str(row_values[0])
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        lead = p.add_run(title)
        lead.bold = True
        lead.font.color.rgb = PALETTE["teal"]
        for idx, value in enumerate(row_values[1:], start=1):
            label = headers[idx] if idx < len(headers) else f"Field {idx}"
            detail = doc.add_paragraph()
            detail.paragraph_format.left_indent = Cm(0.45)
            detail.paragraph_format.space_after = Pt(1)
            label_run = detail.add_run(f"{label}: ")
            label_run.bold = True
            label_run.font.color.rgb = RGBColor(75, 88, 97)
            detail.add_run(str(value))


def add_executive_conclusion(doc: Document, report_inputs: dict) -> None:
    doc.add_heading("1. Executive Conclusion", level=1)
    risk = report_inputs.get("water_risk", {})
    stability = report_inputs.get("stability", {})
    terrain = report_inputs.get("terrain", {})

    p = doc.add_paragraph()
    p.add_run("Preliminary feasibility position. ").bold = True
    p.add_run(
        "A canal, small reservoir/dam, or watercourse diversion can only be considered legally feasible "
        "if it passes two gates: first, the water-management and environmental approvals gate; second, "
        "the technical gate showing no unacceptable harm to water quantity, water quality, flood safety, "
        "ecological status, protected areas, or existing lawful users. The project should therefore be "
        "advanced as a conditional feasibility case, not as a construction-ready decision."
    )

    rows = [
        ["Water demand / risk screening", f"{risk.get('high_risk_count', 'TBD')} high-risk demand clusters; {risk.get('moderate_risk_count', 'TBD')} moderate; {risk.get('low_risk_count', 'TBD')} low.", "Use for prioritization only; permit-grade demand and flow records are still needed."],
        ["Terrain", f"Mean elevation {terrain.get('elevation_mean_m', 'TBD')} m; mean slope {terrain.get('slope_mean_deg', 'TBD')} deg.", "Supports routing and reservoir screening; hydraulic calculations still required."],
        ["Ground stability", f"{stability.get('stability_status', 'TBD')} with score {stability.get('stability_score', 'TBD')}; mean vertical velocity {stability.get('v_mean_mm_per_year', 'TBD')} mm/year.", "Promising, but dam/canal design needs geotechnical investigation."],
        ["Legal gate", "Aviz de gospodarire a apelor before works; autorizatie de gospodarire a apelor before operation.", "No physical intervention should start before the applicable approvals are obtained."],
    ]
    add_matrix_table(doc, ["Topic", "Current indication", "Decision meaning"], rows, [4.2, 5.8, 6.2])

    doc.add_heading("Recommended decision", level=2)
    add_bullets(
        doc,
        [
            "Continue only with a formally scoped feasibility study at SF level, prepared or reviewed by certified specialists.",
            "Prefer the least intrusive option: demand reduction and efficient irrigation first, then off-stream storage, then canal, and only then in-stream dam/diversion if no lower-impact alternative works.",
            "Treat any permanent watercourse diversion, dam, riverbed correction, or floodplain intervention as high regulatory risk until Apele Romane and the environmental authority confirm the procedure.",
        ],
    )


def add_project_and_data(doc: Document, report_inputs: dict) -> None:
    doc.add_heading("2. Project Definition and Data Basis", level=1)
    loc = report_inputs.get("location", {})
    doc.add_paragraph(
        "The assumed project is a rural or peri-urban water infrastructure intervention intended to "
        "supply a community or farm through one or more of the following: surface-water intake, canal, "
        "small retention reservoir, dam/weir, pumping station, gravity feed, or local watercourse diversion."
    )

    add_matrix_table(
        doc,
        ["Item", "Assumption / project input", "Required confirmation"],
        [
            ["Beneficiary", "Community, farm association, municipality, or private farm", "Legal identity, land rights, water-use purpose"],
            ["Coordinates", f"AOI center lat {loc.get('lat', 'TBD')}, lon {loc.get('lon', 'TBD')}", "Final route and structures in Stereo 70 and cadastral parcels"],
            ["Study radius", f"{loc.get('size_km', 'TBD')} km preview AOI", "Final basin/sub-basin boundary and affected water bodies"],
            ["Water source", "EuHydro / clipped water lines and polygons, satellite water masks, GloFAS/EFAS discharge where available", "Official hydrological data from ABA/INHGA as required"],
            ["Demand", "Community proxy from impervious/built-up clusters or farm demand input", "Measured population, irrigation area, crop water need, seasonal schedule"],
            ["Ground context", "Copernicus GLO-30 terrain, EGMS vertical motion, SoilGrids where available", "Field survey, geotechnical boreholes, seepage tests"],
        ],
        [3.6, 6.4, 5.8],
    )

    doc.add_heading("Datasets represented in the project", level=2)
    add_bullets(
        doc,
        [
            "EuHydro or local hydrography layers: clipped rivers, canals, lakes, and water polygons.",
            "Copernicus imperviousness/built-up rasters: proxy for community clusters and demand locations.",
            "Copernicus GLO-30 DEM: elevation, slope, routing, basin screening, and floodplain context.",
            "Sentinel-1 and Sentinel-2 water masks: observed water surface support for river-width and wet/dry checks.",
            "GloFAS/EFAS discharge data: preliminary discharge and daily supply proxy, subject to official validation.",
            "EGMS L3 Ortho Vertical: ground-motion screening for reservoir footprints and canal routes.",
            "SoilGrids: preliminary seepage/permeability indicator, not a replacement for site investigation.",
        ],
    )


def add_legal_framework(doc: Document) -> None:
    doc.add_heading("3. Romanian Legal Framework", level=1)
    doc.add_paragraph(
        "Romanian water infrastructure projects are regulated through an integrated water-management and "
        "environmental framework. For a canal, dam, intake, accumulation, riverbed correction, bank work, "
        "crossing, or diversion, the feasibility study must be written around the following legal controls."
    )

    rows = [
        ["Legea apelor nr. 107/1996", "Core water law; defines works on waters or connected to waters, including dams, accumulations, hydraulic diversions, irrigation water uses, riverbed/bank works, crossings, and similar infrastructure.", "Project is likely within art. 48 categories and should be treated as requiring water-management regulation."],
        ["Art. 50, Legea apelor", "Works in art. 48 may be promoted/executed only on the basis of aviz de gospodarire a apelor or notification, and operation requires autorizatie de gospodarire a apelor where applicable.", "No works or operation before the relevant act is issued."],
        ["Art. 49, Legea apelor", "New objectives in floodable major riverbeds and protection zones are restricted, with limited exceptions subject to flood-defense works and aviz de amplasament.", "Any floodplain location is a major red-flag item."],
        ["Ordin MAP nr. 828/2019", "Procedure for aviz de gospodarire a apelor, including impact assessment on water bodies and technical documentation content.", "The SF-level document must include hydrology, water bodies, project description, monitoring, alternatives, and impacts."],
        ["Ordin MMAP nr. 3147/2023", "Procedure for issuing autorizatie de gospodarire a apelor.", "Relevant after construction, before operation/exploitation."],
        ["Legea nr. 292/2018", "Environmental impact assessment procedure, coordinated by the competent environmental authority and integrated with water-body impact assessment where required.", "Environmental screening and possibly EIA/appropriate assessment must be coordinated early."],
        ["EU Water Framework Directive principles", "No deterioration, achievement of good status/potential, mitigation of hydromorphological alteration, and strict justification for exemptions.", "A project that worsens water-body status is legally difficult and needs exceptional justification."],
    ]
    add_matrix_table(doc, ["Instrument", "Relevance", "How it affects feasibility"], rows, [4.1, 6.2, 5.5])

    doc.add_heading("Legal principles that should guide the study", level=2)
    add_bullets(
        doc,
        [
            "Prior authorization: do not construct, modify the riverbed, dam, divert, abstract, or operate before the required acts are issued.",
            "No deterioration: the project must avoid deterioration of surface-water and groundwater bodies, including hydromorphology and ecological potential.",
            "Sustainable and rational use: abstraction must be proportional to proven need and available resource.",
            "Protection of existing users: downstream users, riparian users, public water supply, ecosystems, and existing permitted works must not be unlawfully impaired.",
            "Flood safety: the project must not increase flood risk upstream, downstream, or on neighboring properties.",
            "Least-damaging feasible alternative: lower-impact options should be analyzed before proposing a dam or permanent diversion.",
            "Polluter-pays and prevention: risks of sediment, turbidity, fuel spills, erosion, stagnation, or water-quality degradation must be prevented rather than remediated later.",
            "Public participation and transparency: environmental procedures may require public information/consultation, especially for projects with significant effects.",
        ],
    )


def add_permitting_roadmap(doc: Document) -> None:
    doc.add_heading("4. Apele Romane Permitting Roadmap", level=1)
    doc.add_paragraph(
        "The feasibility document should be organized so that it can later support the water-management "
        "endorsement procedure. The roadmap below separates the planning, construction, and operating gates."
    )

    rows = [
        ["0", "Pre-screening", "Identify ABA/SGA, affected water body, protected areas, floodplain, land rights, and whether the work falls under art. 48.", "Internal feasibility note, maps, initial authority discussion"],
        ["1", "Environmental initial evaluation", "Submit to the competent environmental authority for the initial evaluation stage.", "Decizia etapei de evaluare initiala or further EIA requirements"],
        ["2", "Urbanism certificate", "Obtain Certificat de urbanism listing required approvals and constraints.", "Certificat de urbanism"],
        ["3", "Water-management endorsement", "Submit request and technical documentation to Apele Romane / ABA / SGA.", "Aviz de gospodarire a apelor and, where relevant, aviz de amplasament"],
        ["4", "Other specialist approvals", "Transport/navigation, forestry, protected areas, land reclamation, dam safety, public health, utilities, depending on site.", "Specialist avize/acorduri"],
        ["5", "Construction authorization", "Apply to local authority after required endorsements.", "Autorizatie de construire"],
        ["6", "Execution and monitoring", "Build according to approved design; announce start where required; record as-built conditions.", "Construction records, monitoring data, reception documents"],
        ["7", "Operation authorization", "After completion and technical verification, apply for operation/exploitation authorization.", "Autorizatie de gospodarire a apelor; operating regulation"],
    ]
    add_matrix_table(doc, ["Step", "Gate", "Main action", "Expected output"], rows, [1.2, 3.3, 6.2, 5.1])

    doc.add_heading("Submission package for aviz de gospodarire a apelor", level=2)
    add_bullets(
        doc,
        [
            "Application form required by the applicable order/procedure.",
            "Environmental initial evaluation decision issued by the environmental authority.",
            "Certificat de urbanism.",
            "Documents proving rights over land occupied by works, including any state public-domain land administered by A.N. Apele Romane.",
            "Technical documentation prepared by an attested public or private entity, according to Ordin 828/2019 and authority requirements.",
            "Hydrological, hydraulic, hydrogeological, geotechnical, inundability, and water-management studies as required by project type.",
            "Maps, plans, profiles, cross-sections, coordinates, affected water-body names/codes, and relation to existing water works.",
        ],
    )


def add_baseline(doc: Document, report_inputs: dict) -> None:
    doc.add_heading("5. Baseline Technical Conditions", level=1)
    terrain = report_inputs.get("terrain", {})
    stability = report_inputs.get("stability", {})
    soil = report_inputs.get("soil", {})
    risk = report_inputs.get("water_risk", {})

    rows = [
        ["Elevation min / max", f"{terrain.get('elevation_min_m', 'TBD')} m / {terrain.get('elevation_max_m', 'TBD')} m", "Check DEM voids and water-surface artifacts before final design."],
        ["Mean elevation", f"{terrain.get('elevation_mean_m', 'TBD')} m", "Useful for regional context only."],
        ["Mean / max slope", f"{terrain.get('slope_mean_deg', 'TBD')} deg / {terrain.get('slope_max_deg', 'TBD')} deg", "Steep zones raise erosion, slope-stability, and canal lining concerns."],
        ["Water-risk distribution", f"{risk.get('risk_counts', {})}", "Prioritize high and moderate demand clusters for alternatives screening."],
        ["EGMS points in AOI", f"{stability.get('measurement_points_in_aoi', 'TBD')}", "Remote screening; absence of points along a canal does not prove stability."],
        ["Reservoir motion", f"{stability.get('reservoir_v_mean_mm_per_year', 'TBD')} mm/year; {stability.get('stability_status', 'TBD')}", "Promising where stable, but site investigations remain mandatory."],
        ["Soil screening", soil.get("error") or f"ksat / seepage class available at point: {soil}", "SoilGrids is only preliminary; use boreholes and permeability tests."],
    ]
    add_matrix_table(doc, ["Baseline item", "Current project value", "Interpretation"], rows, [4.0, 5.7, 6.1])

    doc.add_heading("Baseline studies required for a permit-grade document", level=2)
    add_bullets(
        doc,
        [
            "Affected water bodies: official code, name, status, objectives, pressures, and protected zones from the relevant ABA.",
            "Hydrology: long-term flow series, low-flow statistics, flood flows, seasonal availability, and climate/drought sensitivity.",
            "Hydraulics: water levels, velocities, backwater effects, floodplain connectivity, intake capacity, and canal conveyance.",
            "Sediment and morphology: erosion, deposition, bed stability, bank stability, sediment continuity, and maintenance dredging risk.",
            "Hydrogeology: groundwater level, spring/wetland connections, seepage losses, and impact on wells.",
            "Geotechnics: foundation conditions, slope stability, seepage through embankments, liquefaction/settlement where relevant.",
            "Ecology: fish passage, habitats, riparian corridor, Natura 2000/protected areas, minimum/ecological flow requirements.",
            "Water quality: turbidity, nutrients, temperature, stagnation risk, construction pollution, and operational discharge quality.",
        ],
    )


def add_alternatives(doc: Document) -> None:
    doc.add_heading("6. Alternatives Analysis", level=1)
    doc.add_paragraph(
        "A credible feasibility paper should show that the selected option is not simply convenient, but "
        "legally and environmentally preferable to lower-impact alternatives."
    )
    rows = [
        ["A0: No project", "No canal/dam/diversion; manage demand only", "Lowest environmental impact; may fail water-supply objective", "Baseline comparator"],
        ["A1: Demand reduction", "Drip irrigation, scheduling, leakage reduction, storage tanks", "Often cheapest and least regulated", "Must be tested first"],
        ["A2: Off-stream reservoir", "Pump or gravity-feed into storage outside the active river channel", "Less hydromorphological impact than in-stream dam", "Preferred where land is available"],
        ["A3: Lined canal from permitted intake", "Canal route with controlled abstraction and metering", "Feasible if ecological flow and flood safety remain acceptable", "Medium risk"],
        ["A4: Small weir/dam", "Impoundment or diversion structure in watercourse", "May alter continuity, sediment, fish passage, flood risk", "High regulatory risk"],
        ["A5: Permanent watercourse diversion", "Move or redirect the channel", "Major hydromorphological alteration and downstream risk", "Highest regulatory risk"],
        ["A6: Groundwater supply", "Wells or managed aquifer recharge", "May reduce surface impacts but can affect aquifers/wells", "Needs hydrogeological proof"],
    ]
    add_matrix_table(doc, ["Alternative", "Description", "Strength", "Feasibility signal"], rows, [3.2, 4.6, 4.5, 3.3])

    doc.add_heading("Preferred screening logic", level=2)
    add_numbered(
        doc,
        [
            "Confirm demand and seasonal need.",
            "Apply water-saving and non-structural options.",
            "Use off-stream storage or small controlled intake if sufficient.",
            "Use a canal only if route, gravity/pumping, land rights, ecological flow, and maintenance are feasible.",
            "Use a dam or watercourse diversion only if no less harmful feasible alternative exists and the public/community benefit is strong.",
        ],
    )


def add_impact_and_mitigation(doc: Document) -> None:
    doc.add_heading("7. Impact Assessment and Mitigation", level=1)
    rows = [
        ["Reduced downstream flow", "Abstraction or diversion may reduce baseflow and ecological flow.", "Calculate Q95/Q90/low-flow regimes, seasonal abstraction limits, water balance.", "Meter intake; maintain ecological flow; stop abstraction under drought thresholds."],
        ["Hydromorphology", "Dam/weir/canal may alter river continuity, bed profile, sediment, banks.", "Morphological survey, sediment continuity analysis, WFD water-body impact screening.", "Avoid in-channel works; fish pass; sediment bypass; bank-naturalization."],
        ["Flood risk", "Structures may obstruct flow or shift flood levels.", "1D/2D hydraulic modeling for design floods and blockage scenarios.", "Freeboard, overflow routes, spillway, no net flood-level increase."],
        ["Groundwater", "Reservoir/canal can leak or raise/lower local groundwater.", "Boreholes, permeability tests, piezometers, seepage model.", "Lining, cut-off trench, drainage, monitoring wells."],
        ["Water quality", "Construction turbidity, fuel spills, stagnation, thermal changes.", "Construction method statement, water-quality baseline.", "Silt curtains, spill response, staged works, flushing."],
        ["Existing users", "Downstream intakes, wells, irrigation, fishponds, public supply may be affected.", "Inventory of users and permits; stakeholder consultation.", "Priority rules, operating limits, compensation/avoidance where required."],
        ["Protected areas", "Natura 2000 or habitats may trigger appropriate assessment.", "Protected-area screening and ecological field surveys.", "Avoidance, seasonal windows, habitat compensation only where legally acceptable."],
        ["Dam/canal safety", "Failure, seepage, erosion, or differential settlement.", "Geotechnical design, stability analysis, emergency action plan.", "Instrumentation, inspections, alarm plan, maintenance budget."],
    ]
    add_matrix_table(doc, ["Impact", "Risk", "Required analysis", "Mitigation"], rows, [3.1, 4.2, 5.1, 4.6])


def add_compliance(doc: Document) -> None:
    doc.add_heading("8. Compliance Matrix", level=1)
    rows = [
        ["Does the project fall under Legea apelor art. 48?", "Yes, likely, if it includes canal, intake, dam, accumulation, diversion, riverbed/bank work, irrigation use, or crossings.", "Describe each work component and classify it."],
        ["Is aviz de gospodarire a apelor needed before works?", "Yes, likely for art. 48 works unless a notification regime applies.", "Ask competent ABA/SGA during pre-screening."],
        ["Is autorizatie de gospodarire a apelor needed for operation?", "Yes, likely after completion and before exploitation.", "Prepare operating regulation and monitoring plan."],
        ["Is aviz de amplasament needed?", "If in floodable major riverbed, protection zone, or sensitive location.", "Map flood hazard and protection zones."],
        ["Is environmental procedure needed?", "Yes: initial evaluation under the environmental authority; EIA may be required depending on effects.", "Submit project memo and coordinate water-body assessment."],
        ["Can no-deterioration be demonstrated?", "Unknown until water-body status and impact assessment are complete.", "Prepare Ordin 828/2019 impact screening/study where required."],
        ["Are land rights secure?", "Unknown.", "Cadastral plan, ownership/lease/easements, public-domain permissions."],
        ["Are existing users protected?", "Unknown.", "Inventory downstream/upstream users and permitted abstractions/discharges."],
        ["Is dam safety triggered?", "If a dam, dike, permanent/non-permanent retention, or special hydrotechnical work is included.", "Obtain dam-safety classification and required safety approvals."],
    ]
    add_matrix_table(doc, ["Question", "Preliminary answer", "Action before submission"], rows, [5.3, 5.1, 5.2])


def add_documentation_package(doc: Document) -> None:
    doc.add_heading("9. Minimum Documentation Package", level=1)
    add_matrix_table(
        doc,
        ["Document / study", "Purpose", "Responsible party"],
        [
            ["Project memorandum / feasibility note", "Define need, alternatives, location, works, impacts, and permits.", "Beneficiary + designer"],
            ["Topographic and cadastral survey", "Coordinates, property boundaries, route, structures, cross-sections.", "Licensed surveyor"],
            ["Hydrological study", "Flows, low-flow availability, flood flows, water balance.", "Certified hydrologist / INHGA or ABA-confirmed data where required"],
            ["Hydraulic model", "Water levels, backwater, capacity, flood safety.", "Hydraulic engineer"],
            ["Hydrogeological study", "Groundwater and seepage impacts.", "Certified hydrogeologist"],
            ["Geotechnical study", "Foundation, slope stability, seepage, dam/canal stability.", "Geotechnical engineer"],
            ["Water-body impact study", "No-deterioration and WFD objective check.", "Attested specialist under water authority requirements"],
            ["Environmental screening/EIA materials", "Environmental authority procedure, protected areas, public consultation.", "Environmental assessor"],
            ["Operation and maintenance regulation", "Water abstraction, ecological flow, emergency operation, inspection.", "Designer/operator"],
            ["Monitoring plan", "Flow, water quality, groundwater, stability, sediment, ecology.", "Operator + specialists"],
        ],
        [4.8, 6.2, 4.6],
    )

    doc.add_heading("Recommended report annexes", level=2)
    add_bullets(
        doc,
        [
            "General location map and basin map.",
            "Cadastral and land ownership plan.",
            "Water-body map with official codes and protected areas.",
            "Flood hazard/inundability map.",
            "Longitudinal profile and cross-sections for canal and river works.",
            "Reservoir footprint, storage-elevation curve, spillway concept, and dam cross-section if applicable.",
            "Existing water users and infrastructure inventory.",
            "Construction staging and pollution-prevention plan.",
            "Emergency response and dam/canal failure scenario if relevant.",
        ],
    )


def add_conclusions(doc: Document) -> None:
    doc.add_heading("10. Conclusions and Next Actions", level=1)
    doc.add_paragraph(
        "The research conclusion is that a water infrastructure intervention may be feasible, but only "
        "as a conditional project subject to Romanian water-management, environmental, land, and safety "
        "approvals. The project should be framed around public/community need, proportional abstraction, "
        "least-impact alternatives, and no deterioration of affected water bodies."
    )
    add_numbered(
        doc,
        [
            "Choose one pilot location and freeze an initial alignment/footprint for review.",
            "Identify the competent ABA/SGA and request the affected water-body code/name and procedural guidance.",
            "Prepare a short pre-application pack with maps, coordinates, works description, water need, and alternatives.",
            "Obtain environmental initial evaluation and certificat de urbanism.",
            "Commission the hydrological, hydraulic, geotechnical, hydrogeological, and water-body impact studies required for the selected option.",
            "Only then decide whether the project remains a canal, off-stream storage scheme, dam/weir, or no-build/demand-management solution.",
        ],
    )


def add_appendices(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix A. Legal Source Notes", level=1)
    refs = [
        ["A.N. Apele Romane - Avize si Autorizatii", "https://rowater.ro/activitatea-institutiei/departamente/managementul-european-integrat-resurse-de-apa/avize-si-autorizatii/", "Official explanation of aviz/autorizatie and required documentation."],
        ["Legea apelor nr. 107/1996, art. 48", "https://legislatie.just.ro/Public/FormaPrintabila/00000G2GI51V145KG471KWKZJQSC4N9Z", "Categories of works built on waters or connected to waters."],
        ["Legea apelor nr. 107/1996, art. 49", "https://legeaz.net/legea-apelor-107-1996/art-49-gospodarirea-apelor", "Floodplain/protection-zone restrictions and aviz de amplasament context."],
        ["Legea apelor nr. 107/1996, art. 50", "https://legeaz.net/legea-apelor-107-1996/art-50-gospodarirea-apelor", "Aviz before works and authorization before operation for relevant works."],
        ["Ordin MAP nr. 828/2019", "https://legislatie.just.ro/Public/DetaliiDocument/216574", "Procedure and content for aviz de gospodarire a apelor and water-body impact assessment."],
        ["Continut-cadru study under Ordin 828/2019", "https://legislatie.just.ro/Public/DetaliiDocument/216746", "Technical content expectations for study of impact on water bodies."],
        ["Legea nr. 292/2018", "https://lege5.ro/Gratuit/gmytenbvhezq/legea-nr-292-2018-privind-evaluarea-impactului-anumitor-proiecte-publice-si-private-asupra-mediului", "Environmental impact assessment framework; verify latest consolidated form before filing."],
    ]
    add_matrix_table(doc, ["Source", "URL", "Use in report"], refs, [4.2, 6.5, 5.1])

    doc.add_heading("Appendix B. Technical Checklist", level=1)
    checklist = [
        ["Need and demand", "Population/farm demand, seasonal demand, crop water needs, drought scenario", "Open"],
        ["Water source", "Official flow data, ecological flow, abstraction limit, existing users", "Open"],
        ["Route/footprint", "Coordinates, land rights, cadastral parcels, protected zones", "Open"],
        ["Hydraulics", "Capacity, velocities, flood levels, spillway, outlet, erosion", "Open"],
        ["Geotechnics", "Foundation, embankment, seepage, slope stability, settlement", "Open"],
        ["Environment", "Water-body status, habitats, fish, Natura 2000, water quality", "Open"],
        ["Safety", "Dam class/category, emergency plan, alarm, monitoring instrumentation", "Open"],
        ["Operation", "Operating rule curve, metering, seasonal restrictions, maintenance budget", "Open"],
        ["Legal", "Aviz, autorizatie, urbanism, land, environmental decision, specialist approvals", "Open"],
    ]
    add_matrix_table(doc, ["Area", "Evidence required", "Status"], checklist, [3.8, 9.2, 2.2])


def build() -> Path:
    OUTPUT_DIR.mkdir(exist_ok=True)
    report_inputs = load_json(OUTPUT_DIR / "report_inputs.json")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)
    set_document_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "Romanian Water Infrastructure Feasibility Study"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(90, 103, 112)
    footer = section.footer.paragraphs[0]
    footer.text = "Draft feasibility research report"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(90, 103, 112)

    add_cover(doc, report_inputs)
    add_static_toc(doc)
    add_executive_conclusion(doc, report_inputs)
    add_project_and_data(doc, report_inputs)
    add_legal_framework(doc)
    add_permitting_roadmap(doc)
    add_baseline(doc, report_inputs)
    add_alternatives(doc)
    add_impact_and_mitigation(doc)
    add_compliance(doc)
    add_documentation_package(doc)
    add_conclusions(doc)
    add_appendices(doc)

    doc.core_properties.title = "Legal and Technical Feasibility Study for Water Infrastructure in Romania"
    doc.core_properties.subject = "Canal, dam, watercourse diversion, Apele Romane permitting"
    doc.core_properties.keywords = "Apele Romane, aviz de gospodarire a apelor, autorizatie, canal, dam, water diversion, feasibility"
    doc.core_properties.author = "Codex"
    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build())
