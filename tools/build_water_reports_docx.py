from __future__ import annotations

import json
import math
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

try:
    from pyproj import Transformer
except ImportError:
    Transformer = None

try:
    import rasterio
    from rasterio.vrt import WarpedVRT
except ImportError:
    rasterio = None
    WarpedVRT = None


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output"
ASSET_DIR = OUTPUT_DIR / "report_assets"
GUIDELINE_PATH = OUTPUT_DIR / "romania_water_legal_guideline.docx"
CASE_STUDY_PATH = OUTPUT_DIR / "technical_feasibility_case_study_template.docx"
REPORT_INPUTS_PATH = OUTPUT_DIR / "report_inputs.json"
TERRAIN_SUMMARY_PATH = OUTPUT_DIR / "terrain_summary.json"
STABILITY_SUMMARY_PATH = OUTPUT_DIR / "stability_summary.json"
WATER_RISK_SUMMARY_PATH = OUTPUT_DIR / "water_risk_summary.json"
CANALS_PATH = OUTPUT_DIR / "water_risk_canals.geojson"
SITES_PATH = OUTPUT_DIR / "water_risk_sites.geojson"
TERRAIN_DEM_PATH = OUTPUT_DIR / "terrain_dem.tif"


COLORS = {
    "ink": RGBColor(29, 47, 59),
    "teal": RGBColor(34, 105, 112),
    "muted": RGBColor(88, 103, 112),
}

LEGAL_REFERENCES = [
    {
        "rule": "Legea apelor nr. 107/1996",
        "source": "Portal Legislativ",
        "url": "https://legislatie.just.ro/Public/DetaliiDocument/267923",
        "applies": "Core water law for works built on, in, near, or functionally connected to water bodies.",
        "check": "Confirm whether the canal, lake/reservoir, intake, dam, diversion, crossing, or abstraction falls under water-management endorsement/authorization.",
    },
    {
        "rule": "Ordinul nr. 828/2019",
        "source": "Portal Legislativ",
        "url": "https://legislatie.just.ro/Public/DetaliiDocument/216574",
        "applies": "Procedure, competence, technical-documentation content, and water-body impact study content for water-management endorsement.",
        "check": "Build the technical file around hydrology, hydraulics, water-body impact, alternatives, and mitigation before the ABA/SGA submission.",
    },
    {
        "rule": "Legea nr. 292/2018",
        "source": "Portal Legislativ",
        "url": "https://legislatie.just.ro/Public/DetaliiDocumentAfis/208590",
        "applies": "Environmental impact assessment procedure for public/private projects likely to have significant environmental effects.",
        "check": "Run environmental screening before final design; complete EIA where the competent authority requires it.",
    },
    {
        "rule": "OUG nr. 57/2007",
        "source": "Portal Legislativ",
        "url": "https://legislatie.just.ro/Public/DetaliiDocument/100115",
        "applies": "Protected areas and Natura 2000 appropriate assessment where a project may significantly affect a protected site.",
        "check": "Screen the AOI and route/footprint against Natura 2000 and protected-area layers; avoid or assess significant effects.",
    },
    {
        "rule": "Water Framework Directive 2000/60/EC",
        "source": "EUR-Lex",
        "url": "https://eur-lex.europa.eu/legal-content/EN/TXT/?uri=CELEX%3A32000L0060",
        "applies": "EU water-body objectives, including no deterioration and good ecological/chemical status principles.",
        "check": "Show that the selected option does not deteriorate water-body status or justify any lawful exception through the formal procedure.",
    },
    {
        "rule": "Habitats Directive 92/43/EEC",
        "source": "EUR-Lex",
        "url": "https://eur-lex.europa.eu/legal-content/EN/ALL/?uri=CELEX%3A31992L0043",
        "applies": "Natura 2000 conservation and appropriate-assessment principles.",
        "check": "If protected habitats/species can be affected, the project must pass appropriate assessment before approval.",
    },
]


def configure_doc(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.2)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in [
        ("Title", 24, COLORS["ink"]),
        ("Heading 1", 16, COLORS["ink"]),
        ("Heading 2", 12.5, COLORS["teal"]),
        ("Heading 3", 11, COLORS["teal"]),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name in {"Title", "Heading 1"} else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = COLORS["muted"]

    footer = section.footer.paragraphs[0]
    footer.text = f"Generated {date.today().isoformat()} - HeavyWater"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = COLORS["muted"]


def add_title(doc: Document, title: str, subtitle: str, note: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("DRAFT DOCUMENT")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = COLORS["teal"]

    title_p = doc.add_paragraph(style="Title")
    title_p.add_run(title)
    subtitle_run = title_p.add_run(f"\n{subtitle}")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = COLORS["teal"]

    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_before = Pt(12)
    lead = note_p.add_run("Use note: ")
    lead.bold = True
    lead.font.color.rgb = COLORS["ink"]
    note_p.add_run(note)
    doc.add_page_break()


def add_kv(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.space_after = Pt(2)
    lead = p.add_run(f"{label}: ")
    lead.bold = True
    lead.font.color.rgb = COLORS["teal"]
    p.add_run(value)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1.5)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(1.5)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in header_cells[index].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = COLORS["ink"]
        if widths_cm:
            header_cells[index].width = Cm(widths_cm[index])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths_cm:
                cells[index].width = Cm(widths_cm[index])
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph()


def add_record(doc: Document, title: str, fields: list[tuple[str, str]]) -> None:
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(5)
    heading.paragraph_format.space_after = Pt(1)
    run = heading.add_run(title)
    run.bold = True
    run.font.color.rgb = COLORS["teal"]
    for label, value in fields:
        add_kv(doc, label, value)


def build_guideline() -> Path:
    doc = Document()
    configure_doc(doc, "Romania Water Legal Guideline")
    add_title(
        doc,
        "Romania Water Legal Guideline",
        "General legislative guide for canals, reservoirs, dams, intakes, and watercourse diversions",
        "This guideline is intentionally general. It should remain separate from a project case study, because it explains the legal route and principles that apply before site-specific data is introduced.",
    )

    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This document is a general Romanian legislative guideline for early planning of water infrastructure. "
        "It helps a farm, community, municipality, or consultant understand which legal gates usually apply before "
        "promoting a canal, reservoir, dam, intake, or diversion."
    )
    add_bullets(
        doc,
        [
            "Use it before selecting a final project alternative.",
            "Use it to prepare discussions with A.N. Apele Romane, ABA, SGA, the environmental authority, and the local authority.",
            "Do not use it as the final permit dossier or as a substitute for attested technical documentation.",
        ],
    )

    doc.add_heading("2. Main Legal Principles", level=1)
    add_record(
        doc,
        "Prior approval",
        [
            ("Principle", "Works on waters or connected to waters generally require water-management endorsement before execution."),
            ("Practical meaning", "No excavation, riverbed intervention, damming, diversion, or operation should start before the competent authority confirms the applicable procedure."),
        ],
    )
    add_record(
        doc,
        "No deterioration",
        [
            ("Principle", "The project must avoid deterioration of surface-water and groundwater bodies."),
            ("Practical meaning", "The study must assess hydrology, morphology, ecological flow, water quality, protected areas, sediment, and downstream users."),
        ],
    )
    add_record(
        doc,
        "Least-impact alternative",
        [
            ("Principle", "Demand reduction, efficient irrigation, off-stream storage, and controlled abstraction should be checked before more intrusive works."),
            ("Practical meaning", "A dam or permanent diversion needs a stronger justification than a lower-impact canal or off-stream reservoir."),
        ],
    )
    add_record(
        doc,
        "Flood safety",
        [
            ("Principle", "The project must not increase flood risk on neighboring land, settlements, roads, or downstream assets."),
            ("Practical meaning", "Floodplain occupation, backwater effects, bridge/culvert restrictions, and emergency spill paths must be checked."),
        ],
    )

    doc.add_heading("3. Typical Legal Instruments", level=1)
    for title, fields in [
        (
            "Legea apelor nr. 107/1996",
            [
                ("Role", "Core Romanian water law for works built on waters or connected to waters."),
                ("Relevant project types", "Dams, reservoirs, diversions, canals, irrigation water uses, riverbed corrections, bank works, crossings, intakes, and similar works."),
            ],
        ),
        (
            "Aviz de gospodarire a apelor",
            [
                ("Role", "Water-management endorsement normally obtained before works are promoted/executed."),
                ("When used", "Feasibility/design stage, before construction authorization and physical intervention."),
            ],
        ),
        (
            "Autorizatie de gospodarire a apelor",
            [
                ("Role", "Water-management authorization for operation/exploitation."),
                ("When used", "After execution/reception, before operation, abstraction, storage, discharge, or exploitation where applicable."),
            ],
        ),
        (
            "Environmental procedure",
            [
                ("Role", "Initial environmental evaluation and, where required, EIA or appropriate assessment."),
                ("When used", "Before final approval route is fixed, especially for protected areas or significant hydromorphological impact."),
            ],
        ),
    ]:
        add_record(doc, title, fields)

    doc.add_heading("4. Authority Roadmap", level=1)
    add_numbered(
        doc,
        [
            "Identify the competent ABA/SGA and affected water body.",
            "Prepare a pre-screening package: map, coordinates, concept works, water need, alternatives, land context.",
            "Request environmental initial evaluation.",
            "Obtain Certificat de urbanism and listed endorsements.",
            "Prepare attested technical documentation for aviz de gospodarire a apelor.",
            "Resolve environmental, land, protected-area, utility, dam-safety, and specialist approvals.",
            "Apply for construction authorization only after required endorsements are obtained.",
            "After construction, apply for autorizatie de gospodarire a apelor before operation.",
        ],
    )

    doc.add_heading("5. Minimum Technical Studies", level=1)
    add_bullets(
        doc,
        [
            "Hydrological study: long-term flows, low-flow regime, flood flows, drought conditions, ecological/minimum flow.",
            "Hydraulic study: canal capacity, intake behavior, water levels, flood effect, backwater, erosion velocities.",
            "Geotechnical and hydrogeological study: soil composition, permeability, seepage, settlement, slope stability, groundwater interaction.",
            "Water-body impact study: no-deterioration analysis, ecological status/potential, sediment continuity, morphology, water quality.",
            "Environmental study: protected areas, habitats, fish passage, riparian corridor, construction impacts, cumulative impacts.",
            "Operation plan: abstraction limits, drought restrictions, metering, maintenance, inspections, emergency plan.",
        ],
    )

    doc.add_heading("6. Rules to Apply and Bibliography", level=1)
    doc.add_paragraph(
        "Apply these checks before accepting any HeavyWater recommendation as a project candidate. "
        "The URLs point to the legal source or official EU legal text used for the planning rule."
    )
    add_legal_reference_records(doc)

    doc.add_heading("7. Guideline Conclusion", level=1)
    doc.add_paragraph(
        "Use this guideline as the reusable legal checklist, then use the generated report for the selected canal or lake/reservoir candidate."
    )

    doc.core_properties.title = "Romania Water Legal Guideline"
    doc.core_properties.subject = "General Romanian water-management legislative guide"
    doc.save(GUIDELINE_PATH)
    return GUIDELINE_PATH


def load_json(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def load_geojson_features(path: Path) -> list[dict]:
    if not path.exists():
        return []
    payload = json.loads(path.read_text(encoding="utf-8"))
    return payload.get("features", [])


def choose_candidate(canals: list[dict], sites: list[dict]) -> tuple[str, dict | None]:
    best_canal = max(canals, key=lambda item: float((item.get("properties") or {}).get("option_score") or 0.0), default=None)
    best_site = max(sites, key=lambda item: float((item.get("properties") or {}).get("option_score") or 0.0), default=None)
    canal_score = float((best_canal or {}).get("properties", {}).get("option_score") or 0.0) if best_canal else 0.0
    site_score = float((best_site or {}).get("properties", {}).get("option_score") or 0.0) if best_site else 0.0
    if canal_score >= site_score and best_canal:
        return "canal", best_canal
    if best_site:
        return "reservoir", best_site
    return "none", None


def feature_props(feature: dict | None) -> dict:
    return (feature or {}).get("properties") or {}


def build_candidate_rows(features: list[dict], kind: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for index, feature in enumerate(features, start=1):
        props = feature_props(feature)
        if kind == "canal":
            length = props.get("canal_length_m")
            gravity = props.get("gravity_feasibility_pct")
            soil = props.get("route_seepage_class")
            ksat = props.get("route_ksat_mm_per_hour")
            clay = props.get("route_clay_pct")
            sand = props.get("route_sand_pct")
            silt = props.get("route_silt_pct")
        else:
            length = props.get("feed_canal_length_m")
            gravity = props.get("gravity_feasibility_pct")
            soil = props.get("seepage_class")
            ksat = props.get("ksat_mm_per_hour")
            clay = props.get("clay_pct")
            sand = props.get("sand_pct")
            silt = props.get("silt_pct")
        rows.append(
            [
                f"{kind.title()} {index}",
                props.get("decision", "[no decision]"),
                format_number(props.get("option_score"), 1),
                format_number(length, 0, " m"),
                format_number(gravity, 1, "%"),
                format_number(props.get("demand_m3_day"), 1, " m3/day"),
                format_number(props.get("supply_discharge_m3s"), 3, " m3/s"),
                f"{format_number(clay, 1, '%')} / {format_number(sand, 1, '%')} / {format_number(silt, 1, '%')}",
                f"{available_text(soil, 'n/a')} ({format_number(ksat, 2, ' mm/h')})",
            ]
        )
    return rows


def add_candidate_records(doc: Document, features: list[dict], kind: str) -> None:
    if not features:
        doc.add_paragraph(f"No {kind} candidate was generated in the latest run.")
        return
    for index, feature in enumerate(features, start=1):
        props = feature_props(feature)
        if kind == "canal":
            length = props.get("canal_length_m")
            gravity = props.get("gravity_feasibility_pct")
            soil = props.get("route_seepage_class")
            ksat = props.get("route_ksat_mm_per_hour")
            clay = props.get("route_clay_pct")
            sand = props.get("route_sand_pct")
            silt = props.get("route_silt_pct")
            behavior = props.get("route_soil_behavior")
        else:
            length = props.get("feed_canal_length_m")
            gravity = props.get("gravity_feasibility_pct")
            soil = props.get("seepage_class")
            ksat = props.get("ksat_mm_per_hour")
            clay = props.get("clay_pct")
            sand = props.get("sand_pct")
            silt = props.get("silt_pct")
            behavior = props.get("engineering_note")
        add_record(
            doc,
            f"{kind.title()} option {index}",
            [
                ("Decision", available_text(props.get("decision"), "[no decision]")),
                ("Score", format_number(props.get("option_score"), 1)),
                ("Length", format_number(length, 0, " m")),
                ("Gravity feasibility", format_number(gravity, 1, "%")),
                ("Demand", format_number(props.get("demand_m3_day"), 1, " m3/day")),
                ("Flow rate", format_number(props.get("supply_discharge_m3s"), 3, " m3/s")),
                ("Available flow", format_number(props.get("supply_m3_day"), 1, " m3/day")),
                ("Soil composition", f"clay {format_number(clay, 1, '%')}; sand {format_number(sand, 1, '%')}; silt {format_number(silt, 1, '%')}"),
                ("Seepage / Ksat", f"{available_text(soil, 'n/a')}; {format_number(ksat, 2, ' mm/h')}"),
                ("Soil meaning", available_text(behavior, "Soil interpretation unavailable.")),
                ("Reason", available_text(props.get("decision_reason"), "No decision reason was generated.")),
            ],
        )


def add_legal_reference_records(doc: Document, include_url: bool = True) -> None:
    for item in LEGAL_REFERENCES:
        fields = [
            ("Source", item["source"]),
            ("Rule to apply", item["check"]),
            ("Why it matters", item["applies"]),
        ]
        if include_url:
            fields.append(("URL", item["url"]))
        add_record(doc, item["rule"], fields)


def find_demand_row(summary: dict, demand_id: str | None) -> dict:
    rows = summary.get("report_rows", [])
    for row in rows:
        if row.get("demand_id") == demand_id:
            return row
    return {}


def format_number(value, decimals: int = 2, suffix: str = "") -> str:
    if value is None:
        return "[not available]"
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return str(value)
    if math.isnan(numeric) or math.isinf(numeric):
        return "[not available]"
    return f"{numeric:.{decimals}f}{suffix}"


def format_int(value) -> str:
    if value is None:
        return "[not available]"
    try:
        numeric = int(round(float(value)))
    except (TypeError, ValueError):
        return str(value)
    return f"{numeric}"


def available_text(value: str | None, fallback: str) -> str:
    if value is None:
        return fallback
    text = str(value).strip()
    if not text:
        return fallback
    return text


def format_coords(x: float | None, y: float | None) -> str:
    if x is None or y is None:
        return "[coordinates not available]"
    if Transformer is None:
        return f"EPSG:3035 {float(x):.1f}, {float(y):.1f}"
    transformer = Transformer.from_crs("EPSG:3035", "EPSG:4326", always_xy=True)
    lon, lat = transformer.transform(float(x), float(y))
    return f"{lat:.5f}, {lon:.5f}"


def feasibility_rating(score: float | None) -> str:
    if score is None:
        return "[not available]"
    if score >= 75:
        return "Feasible"
    if score >= 60:
        return "Conditionally feasible"
    if score >= 45:
        return "High risk"
    return "Not recommended"


def selected_option_metrics(candidate_kind: str, props: dict) -> dict:
    if candidate_kind == "reservoir":
        return {
            "length": props.get("feed_canal_length_m"),
            "gravity": props.get("gravity_feasibility_pct"),
            "slope": props.get("local_slope_deg"),
            "soil_behavior": props.get("engineering_note"),
            "ksat": props.get("ksat_mm_per_hour"),
            "clay": props.get("clay_pct"),
            "sand": props.get("sand_pct"),
            "silt": props.get("silt_pct"),
        }
    return {
        "length": props.get("canal_length_m"),
        "gravity": props.get("gravity_feasibility_pct"),
        "slope": props.get("mean_route_slope_deg"),
        "soil_behavior": props.get("route_soil_behavior"),
        "ksat": props.get("route_ksat_mm_per_hour"),
        "clay": props.get("route_clay_pct"),
        "sand": props.get("route_sand_pct"),
        "silt": props.get("route_silt_pct"),
    }


def sample_route_profile(route_coords: list[list[float]], dem_path: Path) -> list[tuple[float, float]]:
    if not route_coords or not dem_path.exists() or rasterio is None or WarpedVRT is None:
        return []
    with rasterio.open(dem_path) as src:
        with WarpedVRT(src, crs="EPSG:3035") as vrt:
            profile = []
            total = 0.0
            last = None
            for x, y in route_coords:
                if last is not None:
                    total += math.dist(last, (x, y))
                try:
                    sample = next(vrt.sample([(x, y)]))
                    value = float(sample[0])
                except Exception:
                    value = float("nan")
                if math.isfinite(value):
                    profile.append((total, value))
                last = (x, y)
    if len(profile) <= 24:
        return profile
    step = max(1, len(profile) // 24)
    thinned = profile[::step]
    if thinned[-1] != profile[-1]:
        thinned.append(profile[-1])
    return thinned


def approximate_route_profile(props: dict) -> list[tuple[float, float]]:
    length_m = float(props.get("canal_length_m") or props.get("feed_canal_length_m") or 0.0)
    elevation_drop_m = props.get("elevation_drop_m")
    if length_m <= 0:
        return []
    try:
        total_drop = abs(float(elevation_drop_m)) if elevation_drop_m is not None else max(length_m * 0.015, 12.0)
    except (TypeError, ValueError):
        total_drop = max(length_m * 0.015, 12.0)
    points = []
    normalized_steps = [0.0, 0.12, 0.28, 0.45, 0.63, 0.81, 1.0]
    height_steps = [0.0, 0.06, -0.03, 0.28, 0.46, 0.71, 1.0]
    start_elevation = 100.0
    for fraction, height_fraction in zip(normalized_steps, height_steps):
        distance = length_m * fraction
        elevation = start_elevation + total_drop * height_fraction
        points.append((distance, elevation))
    return points


def make_case_study_diagrams(candidate_kind: str, candidate: dict | None, demand_row: dict, report_inputs: dict) -> tuple[Path, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    route_path = ASSET_DIR / "case_study_route.png"
    profile_path = ASSET_DIR / "case_study_elevation_profile.png"

    font = ImageFont.load_default()
    props = (candidate or {}).get("properties") or {}
    geometry = (candidate or {}).get("geometry") or {}
    route_coords = geometry.get("coordinates") if geometry.get("type") == "LineString" else []
    profile = sample_route_profile(route_coords, TERRAIN_DEM_PATH) if candidate_kind == "canal" else []
    if not profile and candidate_kind == "canal":
        profile = approximate_route_profile(props)
    loc = report_inputs.get("location", {})
    title_text = "Canal route generated from HeavyWater" if candidate_kind == "canal" else "Reservoir footprint generated from HeavyWater"

    img = Image.new("RGB", (1200, 680), "#f5f8f4")
    d = ImageDraw.Draw(img)
    d.rectangle((50, 50, 1150, 630), outline="#9fb3aa", width=3)
    d.text((70, 70), title_text, fill="#1d2f3b", font=font)
    if route_coords:
        xs = [coord[0] for coord in route_coords]
        ys = [coord[1] for coord in route_coords]
        min_x, max_x = min(xs), max(xs)
        min_y, max_y = min(ys), max(ys)
        span_x = max(max_x - min_x, 1.0)
        span_y = max(max_y - min_y, 1.0)
        scaled = []
        for x, y in route_coords:
            px = 120 + ((x - min_x) / span_x) * 860
            py = 560 - ((y - min_y) / span_y) * 360
            scaled.append((px, py))
        d.line(scaled, fill="#1f7cff", width=12)
        d.line(scaled, fill="#76c7ef", width=4)
        sx, sy = scaled[0]
        ex, ey = scaled[-1]
        d.ellipse((sx - 34, sy - 34, sx + 34, sy + 34), fill="#2e7df6", outline="#0b4fb0", width=3)
        d.rectangle((ex - 60, ey - 48, ex + 60, ey + 48), fill="#dbeadf", outline="#3d8f5a", width=3)
        d.text((max(65, sx - 60), min(590, sy + 42)), "Water source", fill="#1d2f3b", font=font)
        d.text((max(65, ex - 75), max(95, ey + 58)), "Demand area", fill="#1d2f3b", font=font)
    else:
        d.text((120, 320), "No canal route is available in the latest output.", fill="#1d2f3b", font=font)
    d.text(
        (150, 125),
        "\n".join(
            [
                f"Length: {format_number(props.get('canal_length_m') or props.get('feed_canal_length_m'), 1, ' m')}",
                f"Region: AOI near {loc.get('lat', '[lat]')}, {loc.get('lon', '[lon]')}",
                f"Selected alternative: {props.get('decision', '[decision pending]')}",
                f"Gravity feasibility: {format_number(props.get('gravity_feasibility_pct'), 1, '%')}",
            ]
        ),
        fill="#226970",
        font=font,
    )
    img.save(route_path)

    img = Image.new("RGB", (1200, 520), "#fbfcfa")
    d = ImageDraw.Draw(img)
    d.text((60, 40), "Longitudinal elevation profile - source to demand", fill="#1d2f3b", font=font)
    d.line((90, 430, 1120, 430), fill="#9fb3aa", width=2)
    d.line((90, 95, 90, 430), fill="#9fb3aa", width=2)
    if profile:
        distances = [item[0] for item in profile]
        elevations = [item[1] for item in profile]
        min_elev, max_elev = min(elevations), max(elevations)
        span_d = max(distances[-1] - distances[0], 1.0)
        span_e = max(max_elev - min_elev, 1.0)
        points = []
        for distance, elevation in profile:
            px = 90 + ((distance - distances[0]) / span_d) * 1030
            py = 410 - ((elevation - min_elev) / span_e) * 260
            points.append((px, py))
        d.line(points, fill="#226970", width=6)
        for point in points:
            d.ellipse((point[0] - 7, point[1] - 7, point[0] + 7, point[1] + 7), fill="#2e7df6")
        highest = max(points, key=lambda item: -item[1])
        lowest = min(points, key=lambda item: item[1])
        d.text((105, 445), "0 m", fill="#566770", font=font)
        d.text((975, 445), f"{format_number(distances[-1], 0, ' m')} chainage", fill="#566770", font=font)
        d.text((18, 95), "Elevation\n[m]", fill="#566770", font=font)
        d.text((points[0][0] + 10, points[0][1] + 12), "Source intake", fill="#1d2f3b", font=font)
        d.text((highest[0] - 30, highest[1] - 34), "High point", fill="#1d2f3b", font=font)
        d.text((points[-1][0] - 70, points[-1][1] - 28), "Delivery point", fill="#1d2f3b", font=font)
        d.text((920, 90), f"Relief span: {format_number(max_elev - min_elev, 1, ' m')}", fill="#226970", font=font)
    else:
        d.text((120, 230), "Elevation profile is unavailable because no sampled route was produced.", fill="#1d2f3b", font=font)
    img.save(profile_path)
    return route_path, profile_path


def build_case_study() -> Path:
    report_inputs = load_json(REPORT_INPUTS_PATH)
    terrain_summary = load_json(TERRAIN_SUMMARY_PATH)
    stability_summary = load_json(STABILITY_SUMMARY_PATH)
    water_risk_summary = load_json(WATER_RISK_SUMMARY_PATH)
    canals = load_geojson_features(CANALS_PATH)
    sites = load_geojson_features(SITES_PATH)
    candidate_kind, candidate = choose_candidate(canals, sites)
    props = (candidate or {}).get("properties") or {}
    selected_metrics = selected_option_metrics(candidate_kind, props)
    demand_id = props.get("demand_id")
    demand_row = find_demand_row(water_risk_summary, demand_id)
    route_path, profile_path = make_case_study_diagrams(candidate_kind, candidate, demand_row, report_inputs)
    soil_summary = (report_inputs.get("soil") or {}) if isinstance(report_inputs.get("soil"), dict) else {}

    doc = Document()
    configure_doc(doc, "Technical Feasibility Report")
    add_title(
        doc,
        "Technical Feasibility Report",
        "Latest processed canal or reservoir recommendation generated from HeavyWater outputs",
        "This report is filled from the latest HeavyWater processing outputs. Any field that remains unavailable marks a real gap in the current fetch or analysis chain and should be completed during the formal study stage.",
    )

    doc.add_heading("1. Case Study Summary", level=1)
    add_record(
        doc,
        "Recommended alternative",
        [
            ("Selected option", available_text(props.get("decision"), "No candidate selected in the latest run")),
            ("Selected geometry type", "Canal" if candidate_kind == "canal" else "Lake / reservoir" if candidate_kind == "reservoir" else "No geometry selected"),
            ("Region", f"AOI near {report_inputs.get('location', {}).get('lat', '[lat]')}, {report_inputs.get('location', {}).get('lon', '[lon]')}"),
            ("Primary purpose", "Community water supply" if water_risk_summary.get("mode") == "community" else "Farm water supply"),
            ("Feasibility rating", feasibility_rating(props.get("option_score"))),
            ("Main reason", available_text(props.get("decision_reason"), "The latest run did not provide a narrative reason.")),
        ],
    )

    doc.add_heading("2. Input Data Used by the Program", level=1)
    add_bullets(
        doc,
        [
            f"Terrain: DEM available = {'yes' if TERRAIN_DEM_PATH.exists() else 'no'}; elevation mean {format_number(terrain_summary.get('elevation_mean_m'), 1, ' m')}; slope mean {format_number(terrain_summary.get('slope_mean_deg'), 1, ' deg')}.",
            f"Water source: nearest source type {demand_row.get('nearest_source_type', '[not available]')}; source distance {format_number(demand_row.get('distance_to_source_m'), 1, ' m')}; supply {format_number(demand_row.get('supply_m3_day'), 1, ' m3/day')}.",
            f"Soil: route seepage class {available_text(props.get('route_seepage_class') or props.get('seepage_class'), 'not derived in the latest run')}; AOI center soil summary {'available' if soil_summary and not soil_summary.get('error') else 'incomplete in the latest fetch'}.",
            f"Ground movement: status {props.get('canal_stability_status') or props.get('stability_status') or stability_summary.get('stability_status', '[not available]')}; mean movement {format_number(props.get('canal_v_mean_mm_per_year') or props.get('stability_velocity_mm_per_year') or stability_summary.get('v_mean_mm_per_year'), 2, ' mm/year')}.",
            f"Demand: mode {water_risk_summary.get('mode', '[mode]')}; estimated demand {format_number(demand_row.get('demand_m3_day'), 1, ' m3/day')}; population proxy {format_int(demand_row.get('demand_population_proxy'))}.",
            "Constraints: protected areas, cadastral/parcel constraints, and existing users are not yet fully extracted into the report and should still be confirmed in formal studies.",
        ],
    )

    doc.add_heading("3. Candidate Options Generated", level=1)
    if canals or sites:
        add_candidate_records(doc, canals, "canal")
        add_candidate_records(doc, sites, "lake")
    else:
        doc.add_paragraph("No canal or lake/reservoir candidate was generated in the latest run.")

    doc.add_heading("4. Proposed Geometry", level=1)
    add_record(
        doc,
        "Canal option",
        [
            ("Canal length", format_number(props.get("canal_length_m"), 1, " m")),
            ("Start point", format_coords(demand_row.get("supply_sample_x"), demand_row.get("supply_sample_y"))),
            ("End point", f"Demand cluster {demand_id or '[id]'} near {format_coords((candidate or {}).get('geometry', {}).get('coordinates', [[None, None]])[-1][0] if candidate_kind == 'canal' and (candidate or {}).get('geometry', {}).get('coordinates') else None, (candidate or {}).get('geometry', {}).get('coordinates', [[None, None]])[-1][1] if candidate_kind == 'canal' and (candidate or {}).get('geometry', {}).get('coordinates') else None)}"),
            ("Average longitudinal slope", format_number(props.get("mean_route_slope_deg"), 2, " deg")),
            ("Gravity feasibility", format_number(props.get("gravity_feasibility_pct"), 1, "%")),
            ("Crossings", "Crossing inventory not computed in the latest run; derive from parcel, road, and stream intersections before design freeze."),
            ("Earthworks", available_text(props.get("terrain_behavior"), "Earthworks estimate was not derived in the latest run.")),
        ],
    )
    add_record(
        doc,
        "Reservoir option",
        [
            ("Footprint", "No reservoir candidate was selected in the latest run." if candidate_kind != "reservoir" else "Selected lake/reservoir polygon is visible on the map and stored in the latest GeoJSON output."),
            ("Storage volume", "Storage model not yet implemented in the current processing chain."),
            ("Maximum water depth", format_number(props.get("basin_depth_m"), 1, " m")),
            ("Embankment/dam height", "Dam/embankment height not derived because no reservoir geometry was selected."),
            ("Spillway concept", "Spillway type should be fixed during hydraulic design once a reservoir option is selected."),
            ("Seepage control", available_text(props.get("engineering_note") or props.get("route_soil_behavior"), "Seepage-control recommendation is not yet available from the fetched soil data.")),
        ],
    )

    doc.add_picture(str(route_path), width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 1. Generated plan-view representation of the selected conveyance corridor based on the latest HeavyWater candidate geometry.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(8.5)
    cap.runs[0].font.italic = True

    doc.add_picture(str(profile_path), width=Inches(6.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 2. Generated longitudinal elevation profile from the selected route, using sampled terrain where available and metric-derived fallback profiling otherwise.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(8.5)
    cap.runs[0].font.italic = True

    doc.add_heading("5. Water Source Impact", level=1)
    add_record(
        doc,
        "Source balance",
        [
            ("Available flow", f"{format_number(demand_row.get('supply_discharge_m3s'), 3, ' m3/s')} / {format_number(demand_row.get('supply_m3_day'), 1, ' m3/day')} from {demand_row.get('supply_source', '[source]')}"),
            ("Required abstraction", format_number(demand_row.get("demand_m3_day"), 1, " m3/day")),
            ("Ecological/minimum flow", "[still required from formal hydrological study]"),
            ("Impact rating", demand_row.get("water_risk", "[not available]")),
            ("Drought rule", "Reduce or stop abstraction when source supply falls below reserved ecological flow and project demand threshold."),
        ],
    )
    add_bullets(
        doc,
        [
            "Downstream flow reduction must be quantified for normal, dry, and drought scenarios.",
            "Existing users upstream and downstream must be inventoried before the recommendation is accepted.",
            "If modeled supply is lower than demand or source flow is unreliable, the option should be flagged as high legal and technical risk.",
        ],
    )

    doc.add_heading("6. Terrain, Soil, and Movement Findings", level=1)
    add_record(
        doc,
        "Terrain",
        [
            ("Elevation range", f"{format_number(terrain_summary.get('elevation_min_m'), 1, ' m')} to {format_number(terrain_summary.get('elevation_max_m'), 1, ' m')}"),
            ("Mean route slope", format_number(props.get("mean_route_slope_deg") or terrain_summary.get("slope_mean_deg"), 2, " deg")),
            ("Problem zones", available_text(props.get("terrain_behavior"), "Specific terrain problem zones were not separately classified in the latest run.")),
        ],
    )
    add_record(
        doc,
        "Soil composition",
        [
            ("Clay", format_number(selected_metrics.get("clay") or soil_summary.get("clay_pct"), 1, "%")),
            ("Sand", format_number(selected_metrics.get("sand") or soil_summary.get("sand_pct"), 1, "%")),
            ("Silt", format_number(selected_metrics.get("silt") or soil_summary.get("silt_pct"), 1, "%")),
            ("Organic matter", format_number(soil_summary.get("organic_matter_pct"), 1, "%")),
            ("Ksat/permeability", format_number(selected_metrics.get("ksat") or soil_summary.get("ksat_mm_per_hour"), 2, " mm/h")),
            ("Engineering meaning", available_text(selected_metrics.get("soil_behavior") or soil_summary.get("engineering_note") or soil_summary.get("error"), "Soil interpretation was not completed in the latest run.")),
        ],
    )
    add_record(
        doc,
        "Ground movement",
        [
            ("Mean velocity", format_number(props.get("canal_v_mean_mm_per_year") or props.get("stability_velocity_mm_per_year") or stability_summary.get("v_mean_mm_per_year"), 2, " mm/year")),
            ("Differential movement", format_number(stability_summary.get("differential_motion_mm_per_year"), 2, " mm/year")),
            ("Status", props.get("canal_stability_status") or props.get("stability_status") or stability_summary.get("stability_status", "[not available]")),
            ("Design implication", stability_summary.get("maintenance_note") or "Use flexible joints, inspection points, and monitoring if route-specific stability values remain sparse."),
        ],
    )

    doc.add_heading("7. Community or Farm Impact", level=1)
    add_bullets(
        doc,
        [
            f"Expected benefit: demand cluster {demand_id or '[id]'} has estimated need {format_number(demand_row.get('demand_m3_day'), 1, ' m3/day')} and population proxy {format_int(demand_row.get('demand_population_proxy'))}.",
            f"Land impact: built-up area proxy {format_number(demand_row.get('area_m2'), 0, ' m2')} and merged block area {format_number(demand_row.get('block_area_m2'), 0, ' m2')}; parcel/easement extraction is still pending.",
            "Construction impact: temporary access, excavation, turbidity, and traffic remain to be assessed in the formal works method statement.",
            "Operational impact: abstraction control, maintenance access, seasonal restrictions, and metering should be expected for any approved option.",
            f"Residual risks: {demand_row.get('risk_reason', '[risk explanation pending]')}",
        ],
    )

    doc.add_heading("8. Legal Applicability Checklist", level=1)
    add_legal_reference_records(doc, include_url=False)

    doc.add_heading("9. Feasibility Decision", level=1)
    add_record(
        doc,
        "Decision statement",
        [
            ("Result", props.get("decision", "[no decision generated]")),
            ("Conditions", "Proceed only with hydrological, hydraulic, geotechnical, hydrogeological, and legal validation for the selected option."),
            ("Next step", "Prepare a pre-application pack for ABA/SGA and the environmental authority using this route/site as the working candidate."),
        ],
    )

    doc.core_properties.title = "Technical Feasibility Report"
    doc.core_properties.subject = "Canal or reservoir feasibility generated from HeavyWater outputs"
    doc.save(CASE_STUDY_PATH)
    return CASE_STUDY_PATH


def build_all() -> tuple[Path, Path]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return build_guideline(), build_case_study()


if __name__ == "__main__":
    guideline, case_study = build_all()
    print(guideline)
    print(case_study)
